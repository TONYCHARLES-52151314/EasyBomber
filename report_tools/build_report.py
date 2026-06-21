from pathlib import Path
import zlib

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "学号_姓名_EasyBomber.docx"
DIAGRAM = ROOT / "report_tools" / "class_relationship.png"

NAVY = "17365D"
BLUE = "2E74B5"
LIGHT_BLUE = "EAF2F8"
PALE = "F5F7FA"
GRAY = "666666"
LINE = "AAB7C4"
BLACK = "202020"


def set_font(run, name="SimSun", size=10.5, bold=False, color=BLACK, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="8", style="single"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), style)
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def set_table_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)
            tc_pr = row.cells[idx]._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(int(width / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(sum(widths_cm) / 2.54 * 1440)))
    tbl_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_font(run, size=9, color=GRAY)


def body(doc, text, bold_prefix=None, after=5, indent=True):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_font(r1, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    add_bookmark(p, bookmark_name(text))
    return p


def bookmark_name(text):
    normalized = " ".join(text.split())
    return "toc_" + format(zlib.crc32(normalized.encode("utf-8")) & 0xFFFFFFFF, "08x")


def add_bookmark(paragraph, name):
    bookmark_id = str(zlib.crc32(name.encode("utf-8")) & 0x7FFFFFFF)
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), bookmark_id)
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), bookmark_id)
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_toc_entry(doc, number, label, page, level=1, field_start=False, field_end=False):
    full_text = f"{number}  {label}"
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0 if level == 1 else 0.75)
    p.paragraph_format.right_indent = Cm(0.15)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2 if level == 1 else 0)
    p.paragraph_format.line_spacing = 1.0
    tabs = p.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(16.2), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    if field_start:
        field_run = OxmlElement("w:r")
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        field_run.append(begin)
        p._p.append(field_run)
        instr_run = OxmlElement("w:r")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-2" \\h \\z \\u '
        instr_run.append(instr)
        p._p.append(instr_run)
        sep_run = OxmlElement("w:r")
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        sep_run.append(sep)
        p._p.append(sep_run)

    link = OxmlElement("w:hyperlink")
    link.set(qn("w:anchor"), bookmark_name(full_text))
    link.set(qn("w:history"), "1")

    def linked_run(text, size, bold=False):
        wr = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:eastAsia"), "SimSun")
        fonts.set(qn("w:ascii"), "Times New Roman")
        fonts.set(qn("w:hAnsi"), "Times New Roman")
        rpr.append(fonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size * 2)))
        rpr.append(sz)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        rpr.append(color)
        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "none")
        rpr.append(underline)
        if bold:
            rpr.append(OxmlElement("w:b"))
        wr.append(rpr)
        wt = OxmlElement("w:t")
        wt.text = text
        wr.append(wt)
        return wr

    link.append(linked_run(full_text, 10.5 if level == 1 else 9.5, level == 1))
    tab_run = OxmlElement("w:r")
    tab_run.append(OxmlElement("w:tab"))
    link.append(tab_run)
    link.append(linked_run(str(page), 10 if level == 1 else 9.5, False))
    p._p.append(link)
    if field_end:
        end_run = OxmlElement("w:r")
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        end_run.append(end)
        p._p.append(end_run)
    return p


def page_title(doc, number, title, subtitle=None, toc=True):
    p = doc.add_paragraph()
    if toc:
        p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"{number}  {title}")
    set_font(r, name="SimHei", size=18, bold=True, color="000000")
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(14)
        r2 = p2.add_run(subtitle)
        set_font(r2, name="SimSun", size=9.5, color=GRAY)
    else:
        p.paragraph_format.space_after = Pt(10)
    if toc:
        add_bookmark(p, bookmark_name(f"{number}  {title}"))


def callout(doc, label, text, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [17.0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color="C7D8E8", size="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label + "  ")
    set_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_font(r, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def screenshot_placeholders(doc):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [8.25, 8.25])
    prompts = [
        ("运行截图 1｜初始游戏界面", "请截取：窗口标题、完整地图、玩家、6 个怪物、顶部状态栏和底部操作提示。\n\n插入后裁剪为横向画面，并删除本提示文字。"),
        ("运行截图 2｜炸弹或爆炸瞬间", "请截取：炸弹倒计时数字，或十字形爆炸火焰；最好同时体现砖块阻挡/被破坏与怪物受击。\n\n插入后裁剪为横向画面，并删除本提示文字。"),
    ]
    for cell, (title, prompt) in zip(table.rows[0].cells, prompts):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "F8FAFC")
        set_cell_border(cell, color="8FA5B8", size="10", style="dashed")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(title)
        set_font(r, size=11, bold=True, color=BLUE)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(18)
        for i, line in enumerate(prompt.split("\n")):
            if i:
                p2.add_run().add_break()
            rr = p2.add_run(line)
            set_font(rr, size=9, color=GRAY)
        row = table.rows[0]
        row.height = Cm(6.0)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    cap = doc.add_paragraph("图 1  EasyBomber 运行效果（请插入两张本人实际运行截图）")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(5)
    cap.paragraph_format.space_after = Pt(5)
    set_font(cap.runs[0], size=9, color=GRAY)


def code_placeholder(doc, number, title, source, prompt, height_cm=3.2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"代码片段 {number}｜{title}")
    set_font(r, size=10.5, bold=True, color=NAVY)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [17.0])
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, "F3F5F7")
    set_cell_border(cell, color="9AA7B2", size="8", style="dashed")
    row = table.rows[0]
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    pp = cell.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rr = pp.add_run(f"【请粘贴代码】来源：{source}\n{prompt}\n粘贴后保留等宽字体 Consolas 9 pt，并删除本提示文字。")
    set_font(rr, name="Microsoft YaHei", size=9, color=GRAY)


def add_page_break(doc):
    doc.add_page_break()


def create_class_diagram():
    width, height = 1800, 880
    image = Image.new("RGB", (width, height), "#F5F8FC")
    draw = ImageDraw.Draw(image)
    regular_path = r"C:\Windows\Fonts\msyh.ttc"
    bold_path = r"C:\Windows\Fonts\msyhbd.ttc"
    regular = ImageFont.truetype(regular_path, 28)
    small = ImageFont.truetype(regular_path, 22)
    bold = ImageFont.truetype(bold_path, 34)
    title_font = ImageFont.truetype(bold_path, 42)

    draw.rectangle((0, 0, width, 100), fill="#17365D")
    draw.text((70, 25), "EasyBomber 类继承与核心协作关系", font=title_font, fill="white")

    def box(x, y, w, h, title, detail, border="#4C84C3", fill="#252B36"):
        shadow = (x + 10, y + 12, x + w + 10, y + h + 12)
        draw.rounded_rectangle(shadow, radius=22, fill="#D8DEE8")
        draw.rounded_rectangle((x, y, x + w, y + h), radius=22, fill=fill, outline=border, width=7)
        tw = draw.textbbox((0, 0), title, font=bold)[2]
        draw.text((x + (w - tw) / 2, y + 18), title, font=bold, fill="#FFFFFF")
        if detail:
            dw = draw.textbbox((0, 0), detail, font=small)[2]
            draw.text((x + (w - dw) / 2, y + 72), detail, font=small, fill="#D7DEE9")

    def arrow_line(points, color="#3E86D1", width_px=7, arrow=True, dashed=False):
        if dashed:
            for a, b in zip(points[:-1], points[1:]):
                x1, y1 = a; x2, y2 = b
                steps = 14
                for i in range(0, steps, 2):
                    t1, t2 = i / steps, min((i + 1) / steps, 1)
                    draw.line((x1 + (x2-x1)*t1, y1 + (y2-y1)*t1,
                               x1 + (x2-x1)*t2, y1 + (y2-y1)*t2), fill=color, width=width_px)
        else:
            draw.line(points, fill=color, width=width_px, joint="curve")
        if arrow:
            x2, y2 = points[-1]
            x1, y1 = points[-2]
            import math
            ang = math.atan2(y2-y1, x2-x1)
            size = 20
            p1 = (x2 - size*math.cos(ang-0.55), y2 - size*math.sin(ang-0.55))
            p2 = (x2 - size*math.cos(ang+0.55), y2 - size*math.sin(ang+0.55))
            draw.polygon([(x2, y2), p1, p2], fill=color)

    base = (570, 145, 660, 130)
    box(*base, "GameObject", "抽象基类｜update() / draw()", border="#D5B94C")
    child_x = [70, 500, 930, 1360]
    names = [("Player", "玩家输入与移动"), ("Monster", "巡逻与自动更新"),
             ("Bomb", "引信计时"), ("Explosion", "火焰范围与寿命")]
    junction = (900, 350)
    arrow_line([(900, 275), junction], arrow=False)
    for x, (name, detail) in zip(child_x, names):
        center = x + 170
        arrow_line([junction, (center, 430)], arrow=True)
        box(x, 430, 340, 120, name, detail, border="#55B98A")

    # Cooperation layer: deliberately separated from inheritance.
    box(535, 665, 360, 115, "Game", "流程控制与对象容器", border="#7086D8", fill="#2A3140")
    box(1040, 665, 360, 115, "Map", "地图数据与爆炸传播", border="#7086D8", fill="#2A3140")
    arrow_line([(715, 665), (410, 550)], color="#8293A8", width_px=5, arrow=False, dashed=True)
    arrow_line([(715, 665), (670, 550)], color="#8293A8", width_px=5, arrow=False, dashed=True)
    arrow_line([(715, 665), (1100, 550)], color="#8293A8", width_px=5, arrow=False, dashed=True)
    arrow_line([(715, 665), (1530, 550)], color="#8293A8", width_px=5, arrow=False, dashed=True)
    arrow_line([(895, 722), (1040, 722)], color="#8293A8", width_px=5, arrow=True, dashed=True)

    draw.line((80, 835, 185, 835), fill="#3E86D1", width=7)
    draw.text((205, 818), "实线箭头：继承", font=regular, fill="#33485F")
    for i in range(6):
        draw.line((590+i*20, 835, 600+i*20, 835), fill="#8293A8", width=5)
    draw.text((730, 818), "虚线：组合 / 调用协作（不是继承）", font=regular, fill="#33485F")
    image.save(DIAGRAM, quality=95)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(BLACK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.35
    for i, size, before, after, color in [
        (1, 16, 14, 7, BLUE),
        (2, 13, 10, 5, BLUE),
        (3, 11.5, 7, 4, "1F4D78"),
    ]:
        st = doc.styles[f"Heading {i}"]
        st.font.name = "SimHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string("000000")
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True


def configure_sections(doc):
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.header_distance = Cm(0.9)
        section.footer_distance = Cm(0.8)
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header.paragraph_format.space_after = Pt(0)
        r = header.add_run("EasyBomber｜面向对象程序设计课程项目报告")
        set_font(r, name="SimSun", size=8.5, color="666666")
        set_page_number(section.footer.paragraphs[0])


def build():
    create_class_diagram()
    doc = Document()
    configure_styles(doc)
    configure_sections(doc)

    # Cover
    for _ in range(3):
        doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = kicker.add_run("面向对象程序设计 · 课程项目报告")
    set_font(r, name="SimHei", size=12, bold=True, color="000000")
    kicker.paragraph_format.space_after = Pt(22)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    r = title.add_run("简易炸弹人")
    set_font(r, name="SimHei", size=30, bold=True, color="000000")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(50)
    r = sub.add_run("EasyBomber — 基于 C++17 与 EasyX 的单人小游戏")
    set_font(r, size=14, color=GRAY)
    cover_table = doc.add_table(rows=5, cols=2)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(cover_table, [4.0, 9.5])
    cover_rows = [
        ("项目名称", "EasyBomber（简易炸弹人）"),
        ("学号", "________________________________"),
        ("姓名", "________________________________"),
        ("提交日期", "________________________________"),
        ("讲解视频网址", "________________________________"),
    ]
    for i, (label, value) in enumerate(cover_rows):
        c0, c1 = cover_table.rows[i].cells
        set_cell_shading(c0, LIGHT_BLUE)
        set_cell_shading(c1, "FFFFFF")
        set_cell_border(c0, color="D5DFE8", size="6")
        set_cell_border(c1, color="D5DFE8", size="6")
        c0.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0, p1 = c0.paragraphs[0], c1.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r0, r1 = p0.add_run(label), p1.add_run(value)
        set_font(r0, size=10.5, bold=True, color=NAVY)
        set_font(r1, size=10.5, color=BLACK)
        cover_table.rows[i].height = Cm(1.15)
        cover_table.rows[i].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    note = doc.add_paragraph("说明：请在提交前补全个人信息、视频网址、运行截图和代码片段，并将文件名改为“学号_姓名_项目名称.docx”。")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(24)
    set_font(note.runs[0], size=9, color=GRAY)

    add_page_break(doc)
    # Thesis-style linked TOC with real bookmarks and dot-leader tabs.
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_before = Pt(18)
    toc_title.paragraph_format.space_after = Pt(5)
    rr = toc_title.add_run("目  录")
    set_font(rr, size=24, bold=True, color=NAVY)
    toc_sub = doc.add_paragraph()
    toc_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_sub.paragraph_format.space_after = Pt(28)
    rr = toc_sub.add_run("CONTENTS")
    set_font(rr, size=9, color=GRAY)

    toc_rows = [
        ("1", "项目概述", 3, 1),
        ("1.1", "运行截图", 3, 2),
        ("2", "OOP 设计详解", 4, 1),
        ("2.1", "类继承图", 4, 2),
        ("2.2", "抽象基类 GameObject", 4, 2),
        ("2.3", "派生类职责", 4, 2),
        ("2.4", "Game 与 Map：组合优于继承", 5, 2),
        ("2.5", "虚函数如何体现多态", 5, 2),
        ("2.6", "dynamic_cast 的局部使用", 5, 2),
        ("2.7", "如果不使用 OOP", 5, 2),
        ("3", "现代 C++ 特性运用分析", 6, 1),
        ("3.1", "vector 管理动态集合", 6, 2),
        ("3.2", "范围 for 遍历", 6, 2),
        ("3.3", "unique_ptr 与 make_unique（重点）", 7, 2),
        ("3.4", "其他现代风格", 7, 2),
        ("4", "核心难点与解决方案", 8, 1),
        ("4.1", "难点描述", 8, 2),
        ("4.2", "定位思路", 8, 2),
        ("4.3", "解决方案", 8, 2),
        ("4.4", "结果与可改进点", 8, 2),
        ("5", "编译与运行说明", 9, 1),
        ("5.1", "Visual Studio 运行", 9, 2),
        ("5.2", "命令行构建示例", 9, 2),
        ("6", "总结与心得", 9, 1),
    ]
    for idx, (num, label, page, level) in enumerate(toc_rows):
        add_toc_entry(doc, num, label, page, level,
                      field_start=(idx == 0), field_end=(idx == len(toc_rows) - 1))
    toc_note = doc.add_paragraph("提示：目录文字与页码可单击跳转；插入截图或代码后，请核对页码。")
    toc_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_note.paragraph_format.space_before = Pt(24)
    set_font(toc_note.runs[0], size=8.5, color=GRAY)

    add_page_break(doc)
    page_title(doc, "1", "项目概述", "功能、玩法规则与运行效果")
    body(doc, "EasyBomber 是一个使用 C++17 和 EasyX 图形库实现的单人炸弹人小游戏。玩家在 15×21 的网格地图中移动，通过放置定时炸弹清除可破坏砖块并消灭全部怪物。项目采用 Visual Studio C++ 桌面工程组织，地图、玩家、怪物、炸弹和爆炸效果分别由独立类负责。")
    body(doc, "操作规则：方向键控制玩家在道路格上移动；Space 键在当前位置放置炸弹；R 键重新开始；Esc 键退出。炸弹引信为 2 秒，爆炸从中心向上、下、左、右各传播 2 格；固定墙体会立即阻断火焰，可破坏砖块会被炸毁并阻断其后的传播。玩家接触怪物或被火焰命中则失败，消灭全部怪物则胜利。", bold_prefix="操作规则：")
    callout(doc, "项目特点", "固定时间步保护、双缓冲绘制、基于网格的碰撞检测、动态对象统一管理，以及通过虚函数实现的更新/绘制多态。")
    heading(doc, "1.1 运行截图", 2)
    screenshot_placeholders(doc)

    add_page_break(doc)
    page_title(doc, "2", "OOP 设计详解", "重点：继承关系、职责划分与运行时多态")
    heading(doc, "2.1 类继承图", 2)
    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic.paragraph_format.space_after = Pt(2)
    pic.add_run().add_picture(str(DIAGRAM), width=Cm(16.8))
    cap = doc.add_paragraph("图 2  EasyBomber 核心类继承关系")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(cap.runs[0], size=9, color=GRAY)
    heading(doc, "2.2 抽象基类 GameObject", 2)
    body(doc, "GameObject 抽取了所有动态游戏对象共有的网格位置 pos、活动标记 active，以及 update(int dt, const Map&) 和 draw() 两个行为接口。两个接口均为纯虚函数，因此基类不能被直接实例化，派生类必须给出自己的更新和绘制方式。虚析构函数保证通过基类指针销毁派生对象时能够正确执行完整析构过程。")
    body(doc, "active 与 remove() 采用“标记删除”策略：碰撞检测阶段只把对象标记为失效，统一清理阶段再从容器移除。这样可避免在遍历 vector 时立即删除对象造成迭代位置变化或悬空指针。")
    heading(doc, "2.3 派生类职责", 2)
    role_table = doc.add_table(rows=5, cols=3)
    role_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(role_table, [3.0, 6.0, 7.5])
    headers = ["类", "核心状态", "主要职责"]
    for j, text in enumerate(headers):
        cell = role_table.rows[0].cells[j]
        set_cell_shading(cell, "E3EAF1")
        set_cell_border(cell)
        rr = cell.paragraphs[0].add_run(text)
        set_font(rr, bold=True, color=NAVY)
    rows = [
        ("Player", "继承 pos、active", "响应方向键；检查地图与炸弹占位；绘制玩家贴图"),
        ("Monster", "route、routeId、moveTime", "按固定路径巡逻；按 450 ms 节拍移动；遇阻等待"),
        ("Bomb", "timer、FUSE_TIME", "累计引信时间；提供倒计时显示；到时等待 Game 引爆"),
        ("Explosion", "cells、timer", "保存爆炸波及格；判断命中；450 ms 后自动失效"),
    ]
    for i, row in enumerate(rows, 1):
        for j, text in enumerate(row):
            cell = role_table.rows[i].cells[j]
            set_cell_border(cell)
            if i % 2 == 0:
                set_cell_shading(cell, PALE)
            rr = cell.paragraphs[0].add_run(text)
            set_font(rr, size=9.2)

    add_page_break(doc)
    page_title(doc, "2", "OOP 设计详解（续）", "Game、Map 的协作与多态调用", toc=False)
    heading(doc, "2.4 Game 与 Map：组合优于继承", 2)
    body(doc, "Game 不是某一种游戏对象，而是流程编排者，因此没有继承 GameObject。它组合一个 Map、一个 unique_ptr<Player>，以及一个 vector<unique_ptr<GameObject>> 动态对象集合，按“输入—更新—碰撞—清理—绘制”的顺序推进每一帧。Map 同样不是动态实体，它负责二维地图数据、出生点、道路判定、砖块破坏和爆炸传播计算。")
    body(doc, "这一结构体现了“is-a 用继承，has-a 用组合”的原则：Monster 是一种 GameObject，所以采用继承；Game 拥有地图和对象，所以采用组合。职责边界清楚后，地图规则的修改不会侵入贴图绘制，怪物移动逻辑的调整也不会破坏主循环。")
    heading(doc, "2.5 虚函数如何体现多态", 2)
    body(doc, "objects 容器中的静态类型统一为 unique_ptr<GameObject>，但其中实际保存 Monster、Bomb 和 Explosion。Game 在更新和绘制时只调用 objects[i]->update(...) 与 objects[i]->draw()，运行时会根据对象的真实类型选择对应的重写函数：怪物移动、炸弹计时、爆炸消失，以及三类对象各自的绘制逻辑。主循环不需要用大量 if/switch 判断对象种类，这正是运行时多态。")
    callout(doc, "设计价值", "新增一种动态对象时，只需继承 GameObject 并实现 update/draw，再放入 objects 容器；Game 的通用更新和绘制循环无需改写。")
    heading(doc, "2.6 dynamic_cast 的局部使用", 2)
    body(doc, "通用行为通过虚函数处理；只有“仅属于某一种类型”的业务操作才使用 dynamic_cast。例如引爆阶段需要调用 Bomb::isReady()，碰撞阶段需要调用 Explosion::hits()，统计阶段需要识别 Monster。这样的使用范围较小，避免把所有逻辑都退化为类型判断。若后续扩展，可进一步使用对象类别枚举、事件系统或双分派减少类型转换。")
    heading(doc, "2.7 如果不使用 OOP", 2)
    compare = doc.add_table(rows=4, cols=3)
    compare.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(compare, [3.3, 6.6, 6.6])
    data = [
        ("比较项", "过程式写法", "当前 OOP 写法"),
        ("数据组织", "多组全局数组分别保存位置、计时器、贴图和状态", "状态封装在 Player、Monster、Bomb、Explosion 内"),
        ("更新/绘制", "主循环使用大量类型标志与 if/switch 分支", "基类指针统一调用虚函数，自动分派"),
        ("扩展维护", "新增对象需修改多个数组、分支和清理代码", "新增派生类并接入容器，改动范围较集中"),
    ]
    for i, row in enumerate(data):
        for j, text in enumerate(row):
            cell = compare.rows[i].cells[j]
            set_cell_border(cell)
            if i == 0:
                set_cell_shading(cell, "E3EAF1")
            elif i % 2 == 0:
                set_cell_shading(cell, PALE)
            rr = cell.paragraphs[0].add_run(text)
            set_font(rr, size=9.2, bold=(i == 0), color=(NAVY if i == 0 else BLACK))
    body(doc, "因此，OOP 带来的核心收益不是“类更多”，而是变化被限制在职责明确的模块中。可维护性体现在定位问题更直接，可扩展性体现在新功能不必反复改动主循环，资源安全则由 RAII 和智能指针进一步保障。", after=0)

    add_page_break(doc)
    page_title(doc, "3", "现代 C++ 特性运用分析", "从安全性、可读性与简洁性比较现代写法")
    heading(doc, "3.1 vector 管理动态集合", 2)
    body(doc, "项目使用 vector 保存怪物路线、爆炸覆盖格、地图二维数据和动态对象集合。与固定长度 C 数组相比，vector 能根据 push_back 自动扩容，并通过 size() 提供当前元素数量，适合怪物死亡、炸弹生成和爆炸消失这类运行期变化。传统写法需要预估最大容量、维护计数变量，并在插入删除时手工移动元素，容易越界或出现计数不同步。")
    code_placeholder(doc, 1, "动态对象集合", "Game.hpp", "粘贴 objects 成员声明，并可连同 player 的 unique_ptr 声明一起截取。", 2.7)
    heading(doc, "3.2 范围 for 遍历", 2)
    body(doc, "Explosion::draw() 使用范围 for 逐个取得 cells 中的 GridPosition。代码直接表达“对每个爆炸格绘制贴图”，不需要管理下标，也不会误写循环边界。传统下标循环必须写初始化、条件和自增，并反复使用 cells[i]，信息噪声更大。若元素无需修改，还可进一步写成 const GridPosition&，减少复制并表达只读意图。")
    code_placeholder(doc, 2, "遍历爆炸覆盖格", "Explosion.cpp → Explosion::draw()", "粘贴从 for (GridPosition cell : cells) 开始，到循环右花括号结束的代码。", 2.7)

    add_page_break(doc)
    page_title(doc, "3", "现代 C++ 特性运用分析（续）", "智能指针、RAII 与对象创建", toc=False)
    heading(doc, "3.3 unique_ptr 与 make_unique（重点）", 2)
    callout(doc, "自学特性｜智能指针", "本项目使用 unique_ptr 表达对象的唯一所有权，并使用 make_unique 创建 Player、Monster、Bomb 和 Explosion。", fill="FFF4D6")
    body(doc, "如果使用传统裸指针，Game 需要在 reset()、正常退出、对象死亡和异常路径中手动 delete；任何遗漏都可能造成内存泄漏，重复释放又可能导致程序崩溃。unique_ptr 在离开作用域或从 vector 中擦除时自动析构对象，资源生命周期与容器生命周期绑定，符合 RAII 思想。")
    body(doc, "make_unique 还把“分配内存”和“交给智能指针管理”合并成一个表达式，避免 new 与所有权接管之间出现异常窗口。objects.push_back(make_unique<Monster>(...)) 同时清楚表达：创建一个 Monster，并把它的唯一所有权转移给 objects。")
    code_placeholder(doc, 3, "智能指针创建与所有权转移", "Game.cpp → reset()、putBomb()、explodeBombs()", "任选 2～4 行 make_unique 语句粘贴，例如创建 Player、Monster、Bomb 或 Explosion 的语句。", 3.4)
    heading(doc, "3.4 其他现代风格", 2)
    body(doc, "GridPosition 使用聚合初始化 GridPosition p{row, col}，语义直接且避免字段遗漏；派生类方法使用 override，编译器可以检查函数签名是否真正重写基类虚函数；GameObject 的析构函数写成 virtual ~GameObject() = default，既保证多态析构，又避免无意义的手写空函数。这些写法共同减少样板代码，并把设计意图交给编译器校验。")
    body(doc, "需要说明的是，项目部分循环仍采用 int 下标与显式类型，这是为了保持课堂代码易读。若继续改进，可在不涉及删除的遍历中使用范围 for 或 auto，在索引场景中改用 std::size_t，并用 erase-remove_if 简化失效对象清理。")

    add_page_break(doc)
    page_title(doc, "4", "核心难点与解决方案", "动态对象生命周期、遍历删除与同帧碰撞")
    heading(doc, "4.1 难点描述", 2)
    body(doc, "最棘手的设计问题是：炸弹到时需要消失并立即生成 Explosion；爆炸又可能在同一帧杀死怪物或玩家；随后这些失效对象还要从 vector 中删除。如果在遍历 objects 时直接 erase，后续元素会前移，循环下标可能跳过对象；如果保存了被删除对象的裸指针，还可能形成悬空指针。")
    heading(doc, "4.2 定位思路", 2)
    body(doc, "问题可以通过观察“连续对象未被清理”“爆炸偶尔漏判”或调试器中的容器大小变化定位。重点检查三个时间点：对象何时被标记失效、Explosion 何时加入容器、clearObjects() 何时真正擦除。把每一帧拆成明确阶段后，可以发现碰撞检测与删除操作混在一起是风险来源。")
    heading(doc, "4.3 解决方案", 2)
    sol = doc.add_table(rows=4, cols=2)
    sol.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(sol, [3.4, 13.1])
    items = [
        ("① 延迟删除", "remove() 只把 active 设为 false，不在碰撞或引爆遍历中立即 erase。"),
        ("② 分阶段更新", "先检查上一帧碰撞，再更新对象；引爆后再次检查碰撞，使新火焰同帧生效。"),
        ("③ 统一清理", "clearObjects() 最后执行；删除元素后下标不递增，避免跳过前移元素。"),
        ("④ RAII 托管", "vector<unique_ptr<GameObject>> 在 erase 时自动析构对象，不需要手工 delete。"),
    ]
    for i, (label, text) in enumerate(items):
        for j, value in enumerate((label, text)):
            cell = sol.rows[i].cells[j]
            set_cell_border(cell)
            if i % 2 == 0:
                set_cell_shading(cell, PALE)
            rr = cell.paragraphs[0].add_run(value)
            set_font(rr, size=9.7, bold=(j == 0), color=(NAVY if j == 0 else BLACK))
    heading(doc, "4.4 结果与可改进点", 2)
    body(doc, "采用上述流程后，对象状态变化具有确定顺序：本帧产生的爆炸能够立即参与伤害判定，失效对象不会继续 update/draw，容器清理也不会遗漏元素。进一步改进时，可使用 std::erase_if（C++20）或 erase-remove_if（C++11/17）统一删除逻辑，并为“炸弹—爆炸—碰撞—清理”编写无图形界面的单元测试。")
    code_placeholder(doc, 4, "可选：延迟删除与统一清理", "GameObject.hpp 的 remove()；Game.cpp 的 clearObjects()", "如果版面允许，可粘贴 remove() 和 clearObjects()，用于支撑本节分析；不粘贴也可删除整个框。", 2.4)

    add_page_break(doc)
    page_title(doc, "5", "编译与运行说明", "环境、依赖与构建命令")
    env = doc.add_table(rows=5, cols=2)
    env.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(env, [4.2, 12.3])
    env_data = [
        ("操作系统", "Windows 10/11（x64）"),
        ("开发环境", "Visual Studio 2026，安装“使用 C++ 的桌面开发”工作负载"),
        ("编译器/工具集", "MSVC，Platform Toolset v145"),
        ("语言标准", "C++17（满足 C++11 或更高版本要求）"),
        ("图形依赖", "EasyX；链接 EasyXw.lib，运行目录需包含 assets 文件夹"),
    ]
    for i, (label, value) in enumerate(env_data):
        for j, text in enumerate((label, value)):
            cell = env.rows[i].cells[j]
            set_cell_border(cell)
            if i % 2 == 0:
                set_cell_shading(cell, PALE)
            rr = cell.paragraphs[0].add_run(text)
            set_font(rr, size=9.8, bold=(j == 0), color=(NAVY if j == 0 else BLACK))
    heading(doc, "5.1 Visual Studio 运行", 2)
    body(doc, "安装 EasyX 并完成 Visual Studio 配置后，双击 EasyBomber.sln，选择 Debug | x64 或 Release | x64。按 F5 启动调试，或按 Ctrl+F5 直接运行。工程的生成后事件会把 assets 目录复制到 bin\\Debug 或 bin\\Release，避免运行时找不到贴图。")
    heading(doc, "5.2 命令行构建示例", 2)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [17.0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "20252B")
    set_cell_border(cell, color="20252B")
    p = cell.paragraphs[0]
    r = p.add_run('msbuild EasyBomber.sln /p:Configuration=Release /p:Platform=x64\n.\\bin\\Release\\EasyBomber.exe')
    set_font(r, name="Consolas", size=9.5, color="F1F1F1")
    body(doc, "命令应在“Developer PowerShell for Visual Studio”中、项目根目录下执行。若链接阶段提示找不到 EasyXw.lib，应检查 EasyX 是否安装到当前 MSVC 工具集，并核对项目的附加库目录。", indent=False)
    heading(doc, "6  总结与心得", 1)
    body(doc, "通过 EasyBomber 项目，我把继承、虚函数、多态、组合和资源管理落实到了一个可运行的小游戏中。最直接的收获是理解了“共同接口与不同实现”的价值：GameObject 统一 update/draw 后，主循环能够专注于流程，不必知道每个对象的全部细节。")
    body(doc, "项目仍有不足：怪物只沿固定路线巡逻，关卡数据写在源码中，dynamic_cast 使用较多，贴图加载方式也可以抽成资源管理器。后续可增加道具、连锁爆炸、多关卡配置和简单寻路，并把碰撞与地图规则拆成可单元测试的纯逻辑模块。总体而言，本项目让我体会到，OOP 的重点并非堆叠类，而是让职责、所有权和变化边界变得清楚。")
    callout(doc, "提交前检查", "补全封面；插入至少 2 张运行截图；粘贴 3 处现代 C++ 代码；更新图号或页码；将文件名改为“学号_姓名_EasyBomber.docx”。")

    # Avoid headers/footers on cover while retaining them on body pages.
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.first_page_header.paragraphs[0].text = ""
    section.first_page_footer.paragraphs[0].text = ""

    doc.core_properties.title = "EasyBomber 面向对象程序设计课程项目报告"
    doc.core_properties.subject = "C++17 / EasyX / OOP"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "EasyBomber, C++17, OOP, EasyX"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
